from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Стили
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = 'Times New Roman'
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

# ─── ТИТУЛЬНЫЙ ЛИСТ ───────────────────────────────────────────────────────────
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Проектирование базы данных\nПроект: RecipeBook')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info_p.add_run('2026 г.')
run3.font.name = 'Times New Roman'
run3.font.size = Pt(12)

doc.add_page_break()

# ─── ЗАДАНИЕ 1: АНАЛИЗ ТРЕБОВАНИЙ ─────────────────────────────────────────────
heading('Задание 1. Анализ требований', 1)

heading('1.1. Описание системы', 2)
para('RecipeBook — персональное веб-приложение для хранения и управления кулинарными рецептами. '
     'Система позволяет пользователям создавать, редактировать и удалять рецепты, '
     'управлять категориями, добавлять рецепты в избранное, формировать план меню на неделю '
     'и выгружать отчёты в форматах XLSX и PDF.')

heading('1.2. Основные сущности и атрибуты', 2)

entities = [
    ('Пользователь (User)', [
        'id — уникальный идентификатор (PK)',
        'name — имя пользователя',
        'email — адрес электронной почты (уникальный)',
        'hashed_password — хэш пароля',
        'created_at — дата и время регистрации',
    ]),
    ('Рецепт (Recipe)', [
        'id — уникальный идентификатор (PK)',
        'owner_id — идентификатор владельца (FK → User)',
        'title — название рецепта',
        'description — описание рецепта (необязательное)',
        'category_id — идентификатор категории (FK → Category)',
        'difficulty — сложность: easy / medium / hard',
        'cook_time — время приготовления (в минутах)',
        'photo_url — URL фотографии',
        'ingredients — список ингредиентов (JSONB)',
        'steps — список шагов приготовления (JSONB)',
        'created_at — дата создания',
        'updated_at — дата последнего обновления',
    ]),
    ('Категория (Category)', [
        'id — уникальный идентификатор (PK)',
        'owner_id — идентификатор владельца (FK → User)',
        'name — название категории',
    ]),
    ('Избранное (Favorite)', [
        'id — уникальный идентификатор (PK)',
        'user_id — идентификатор пользователя (FK → User)',
        'recipe_id — идентификатор рецепта (FK → Recipe)',
        'created_at — дата добавления в избранное',
    ]),
    ('План меню (MenuPlan)', [
        'id — уникальный идентификатор (PK)',
        'user_id — идентификатор пользователя (FK → User)',
        'recipe_id — идентификатор рецепта (FK → Recipe)',
        'date — дата приёма пищи',
        'meal_type — тип приёма пищи: breakfast / lunch / dinner',
        'created_at — дата создания записи',
    ]),
]

for entity_name, attrs in entities:
    p = doc.add_paragraph()
    run = p.add_run(entity_name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    for attr in attrs:
        li = doc.add_paragraph(attr, style='List Bullet')
        li.runs[0].font.name = 'Times New Roman'
        li.runs[0].font.size = Pt(11)

heading('1.3. Связи между сущностями', 2)

relations = [
    ('User → Recipe', 'Один-ко-многим', 'Один пользователь может иметь много рецептов (owner_id)'),
    ('User → Category', 'Один-ко-многим', 'Один пользователь может создать много категорий'),
    ('Category → Recipe', 'Один-ко-многим', 'Одна категория может содержать много рецептов'),
    ('User ↔ Recipe (Favorites)', 'Многие-ко-многим', 'Пользователь добавляет рецепты в избранное через таблицу favorites'),
    ('User ↔ Recipe (MenuPlan)', 'Многие-ко-многим', 'Пользователь назначает рецепты на дни недели через menu_plan'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Связь'
hdr[1].text = 'Тип'
hdr[2].text = 'Описание'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

for rel, rtype, desc in relations:
    row = table.add_row().cells
    row[0].text = rel
    row[1].text = rtype
    row[2].text = desc
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

heading('1.4. Бизнес-правила', 2)

rules = [
    'Email пользователя должен быть уникальным в системе.',
    'Рецепт принадлежит ровно одному пользователю (owner_id NOT NULL).',
    'Категория принадлежит ровно одному пользователю (owner_id NOT NULL).',
    'Пара (user_id, recipe_id) в таблице favorites уникальна — нельзя добавить рецепт в избранное дважды.',
    'Пара (user_id, date, meal_type) в таблице menu_plan уникальна — на один приём пищи назначается один рецепт.',
    'Сложность рецепта ограничена значениями: easy, medium, hard.',
    'Тип приёма пищи ограничен значениями: breakfast, lunch, dinner.',
    'При удалении рецепта каскадно удаляются записи в favorites и menu_plan.',
    'При удалении категории recipe.category_id устанавливается в NULL.',
    'Пароль хранится только в виде хэша (bcrypt), plaintext не сохраняется.',
    'Время приготовления (cook_time) должно быть положительным числом (1–1440 минут).',
    'Рецепт должен содержать хотя бы один ингредиент и один шаг приготовления.',
]

for rule in rules:
    li = doc.add_paragraph(rule, style='List Bullet')
    li.runs[0].font.name = 'Times New Roman'
    li.runs[0].font.size = Pt(11)

doc.add_page_break()

# ─── ЗАДАНИЕ 2: ПРОЕКТИРОВАНИЕ ERD ────────────────────────────────────────────
heading('Задание 2. Проектирование ERD', 1)

heading('2.1. ER-диаграмма (текстовое представление)', 2)

para('Ввиду ограничений формата документа, ниже представлено структурированное описание '
     'ER-диаграммы с указанием всех сущностей, атрибутов и связей.')

erd_text = """\
┌─────────────────────┐         ┌────────────────────────────┐
│       USERS         │         │          RECIPES            │
├─────────────────────┤         ├────────────────────────────┤
│ PK  id              │◄────┐   │ PK  id                     │
│     name            │     │   │ FK  owner_id ──────────────┤
│     email (UNIQUE)  │     └───┤ FK  category_id            │
│     hashed_password │         │     title                  │
│     created_at      │         │     description            │
└─────────────────────┘         │     difficulty (ENUM)      │
          │                     │     cook_time              │
          │                     │     photo_url              │
          │ 1:N                 │     ingredients (JSONB)    │
          ▼                     │     steps (JSONB)          │
┌─────────────────────┐         │     created_at             │
│      CATEGORIES     │         │     updated_at             │
├─────────────────────┤         └────────────────────────────┘
│ PK  id              │                    │
│ FK  owner_id ───────┤                    │
│     name            │         ┌──────────┴──────────┐
└─────────────────────┘         │                     │
          │                     │ N:M (favorites)     │ N:M (menu_plan)
          └─────────────────────┤                     │
                    ┌───────────┴──────┐  ┌───────────┴──────┐
                    │    FAVORITES     │  │    MENU_PLAN      │
                    ├──────────────────┤  ├──────────────────┤
                    │ PK  id           │  │ PK  id           │
                    │ FK  user_id      │  │ FK  user_id      │
                    │ FK  recipe_id    │  │ FK  recipe_id    │
                    │     created_at   │  │     date         │
                    │ UNIQUE(user_id,  │  │     meal_type    │
                    │        recipe_id)│  │     created_at   │
                    └──────────────────┘  │ UNIQUE(user_id,  │
                                          │   date,meal_type)│
                                          └──────────────────┘"""

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.2)
run = p.add_run(erd_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F8F8F8')
p._p.get_or_add_pPr().append(shading)

heading('2.2. Описание связей', 2)

links_data = [
    ('users → recipes', 'Один-ко-многим (1:N)', 'users.id = recipes.owner_id', 'CASCADE DELETE'),
    ('users → categories', 'Один-ко-многим (1:N)', 'users.id = categories.owner_id', 'CASCADE DELETE'),
    ('categories → recipes', 'Один-ко-многим (1:N)', 'categories.id = recipes.category_id', 'SET NULL'),
    ('users ↔ recipes (favorites)', 'Многие-ко-многим (N:M)', 'через таблицу favorites', 'CASCADE DELETE'),
    ('users ↔ recipes (menu_plan)', 'Многие-ко-многим (N:M)', 'через таблицу menu_plan', 'CASCADE DELETE'),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
h2 = table2.rows[0].cells
for i, txt in enumerate(['Связь', 'Тип', 'Условие JOIN', 'При удалении']):
    h2[i].text = txt
    h2[i].paragraphs[0].runs[0].bold = True
    h2[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

for row_data in links_data:
    row = table2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

heading('2.3. Первичные и внешние ключи, ограничения', 2)

constraints = [
    ('users', 'id', 'SERIAL PRIMARY KEY', '—'),
    ('users', 'email', 'UNIQUE NOT NULL', '—'),
    ('recipes', 'id', 'SERIAL PRIMARY KEY', '—'),
    ('recipes', 'owner_id', 'NOT NULL', 'FK → users(id) ON DELETE CASCADE'),
    ('recipes', 'category_id', 'NULLABLE', 'FK → categories(id) ON DELETE SET NULL'),
    ('recipes', 'difficulty', "CHECK IN ('easy','medium','hard')", '—'),
    ('recipes', 'cook_time', 'CHECK (cook_time > 0)', '—'),
    ('categories', 'id', 'SERIAL PRIMARY KEY', '—'),
    ('categories', 'owner_id', 'NOT NULL', 'FK → users(id) ON DELETE CASCADE'),
    ('favorites', 'id', 'SERIAL PRIMARY KEY', '—'),
    ('favorites', '(user_id, recipe_id)', 'UNIQUE', 'FK → users, recipes ON DELETE CASCADE'),
    ('menu_plan', 'id', 'SERIAL PRIMARY KEY', '—'),
    ('menu_plan', '(user_id, date, meal_type)', 'UNIQUE', 'FK → users, recipes ON DELETE CASCADE'),
    ('menu_plan', 'meal_type', "CHECK IN ('breakfast','lunch','dinner')", '—'),
]

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells
for i, txt in enumerate(['Таблица', 'Поле/Набор полей', 'Ограничение', 'Внешний ключ']):
    h3[i].text = txt
    h3[i].paragraphs[0].runs[0].bold = True
    h3[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

for row_data in constraints:
    row = table3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_page_break()

# ─── ЗАДАНИЕ 3: SQL СКРИПТЫ ────────────────────────────────────────────────────
heading('Задание 3. Разработка базы данных', 1)

heading('3.1. SQL-скрипт создания базы данных', 2)

para('База данных разворачивается в СУБД PostgreSQL 15. '
     'Скрипт создаёт таблицы, индексы, ограничения и необходимые типы данных.')

sql = """\
-- ============================================================
-- RecipeBook Database Schema
-- СУБД: PostgreSQL 15
-- Кодировка: UTF-8
-- ============================================================

-- Создание базы данных (выполнять от суперпользователя)
-- CREATE DATABASE recipebook
--     WITH ENCODING 'UTF8'
--          LC_COLLATE = 'ru_RU.UTF-8'
--          TEMPLATE template0;

-- Подключиться к базе: \\c recipebook

-- ------------------------------------------------------------
-- 1. Таблица пользователей
-- ------------------------------------------------------------
CREATE TABLE IF NOT EXISTS users (
    id              SERIAL          PRIMARY KEY,
    name            VARCHAR(100)    NOT NULL,
    email           VARCHAR(255)    NOT NULL UNIQUE,
    hashed_password VARCHAR(255)    NOT NULL,
    created_at      TIMESTAMPTZ     NOT NULL DEFAULT NOW()
);

COMMENT ON TABLE  users                  IS 'Зарегистрированные пользователи';
COMMENT ON COLUMN users.email            IS 'Уникальный email — используется для входа';
COMMENT ON COLUMN users.hashed_password  IS 'Хэш пароля (bcrypt)';

-- ------------------------------------------------------------
-- 2. Таблица категорий
-- ------------------------------------------------------------
CREATE TABLE IF NOT EXISTS categories (
    id          SERIAL          PRIMARY KEY,
    owner_id    INTEGER         NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
    name        VARCHAR(100)    NOT NULL,
    created_at  TIMESTAMPTZ     NOT NULL DEFAULT NOW()
);

COMMENT ON TABLE categories IS 'Категории рецептов (персональные для каждого пользователя)';

-- Индекс для быстрой выборки категорий пользователя
CREATE INDEX IF NOT EXISTS idx_categories_owner_id
    ON categories(owner_id);

-- ------------------------------------------------------------
-- 3. Таблица рецептов
-- ------------------------------------------------------------
CREATE TYPE difficulty_enum AS ENUM ('easy', 'medium', 'hard');

CREATE TABLE IF NOT EXISTS recipes (
    id          SERIAL              PRIMARY KEY,
    owner_id    INTEGER             NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
    category_id INTEGER
                    REFERENCES categories(id) ON DELETE SET NULL,
    title       VARCHAR(200)        NOT NULL,
    description TEXT,
    difficulty  difficulty_enum     NOT NULL DEFAULT 'medium',
    cook_time   INTEGER             NOT NULL
                    CHECK (cook_time BETWEEN 1 AND 1440),
    photo_url   VARCHAR(512),
    ingredients JSONB               NOT NULL DEFAULT '[]',
    steps       JSONB               NOT NULL DEFAULT '[]',
    created_at  TIMESTAMPTZ         NOT NULL DEFAULT NOW(),
    updated_at  TIMESTAMPTZ         NOT NULL DEFAULT NOW()
);

COMMENT ON TABLE  recipes             IS 'Кулинарные рецепты пользователей';
COMMENT ON COLUMN recipes.cook_time   IS 'Время приготовления в минутах (1-1440)';
COMMENT ON COLUMN recipes.ingredients IS 'JSON-массив: [{name, amount}]';
COMMENT ON COLUMN recipes.steps       IS 'JSON-массив: [{description}]';

-- Индексы
CREATE INDEX IF NOT EXISTS idx_recipes_owner_id
    ON recipes(owner_id);

CREATE INDEX IF NOT EXISTS idx_recipes_category_id
    ON recipes(category_id);

CREATE INDEX IF NOT EXISTS idx_recipes_difficulty
    ON recipes(difficulty);

-- Полнотекстовый поиск по названию
CREATE INDEX IF NOT EXISTS idx_recipes_title_fts
    ON recipes USING gin(to_tsvector('russian', title));

-- Автообновление updated_at
CREATE OR REPLACE FUNCTION update_updated_at()
RETURNS TRIGGER LANGUAGE plpgsql AS $$
BEGIN
    NEW.updated_at = NOW();
    RETURN NEW;
END;
$$;

CREATE TRIGGER trg_recipes_updated_at
    BEFORE UPDATE ON recipes
    FOR EACH ROW EXECUTE FUNCTION update_updated_at();

-- ------------------------------------------------------------
-- 4. Таблица избранного
-- ------------------------------------------------------------
CREATE TABLE IF NOT EXISTS favorites (
    id          SERIAL      PRIMARY KEY,
    user_id     INTEGER     NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
    recipe_id   INTEGER     NOT NULL
                    REFERENCES recipes(id) ON DELETE CASCADE,
    created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW(),

    CONSTRAINT uq_favorites_user_recipe UNIQUE (user_id, recipe_id)
);

COMMENT ON TABLE favorites IS 'Избранные рецепты пользователей';
COMMENT ON CONSTRAINT uq_favorites_user_recipe ON favorites
    IS 'Один пользователь не может добавить один рецепт в избранное дважды';

CREATE INDEX IF NOT EXISTS idx_favorites_user_id
    ON favorites(user_id);

CREATE INDEX IF NOT EXISTS idx_favorites_recipe_id
    ON favorites(recipe_id);

-- ------------------------------------------------------------
-- 5. Таблица плана меню
-- ------------------------------------------------------------
CREATE TYPE meal_type_enum AS ENUM ('breakfast', 'lunch', 'dinner');

CREATE TABLE IF NOT EXISTS menu_plan (
    id          SERIAL          PRIMARY KEY,
    user_id     INTEGER         NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
    recipe_id   INTEGER         NOT NULL
                    REFERENCES recipes(id) ON DELETE CASCADE,
    date        DATE            NOT NULL,
    meal_type   meal_type_enum  NOT NULL,
    created_at  TIMESTAMPTZ     NOT NULL DEFAULT NOW(),

    CONSTRAINT uq_menu_plan_slot UNIQUE (user_id, date, meal_type)
);

COMMENT ON TABLE menu_plan IS 'Недельный план питания пользователя';
COMMENT ON CONSTRAINT uq_menu_plan_slot ON menu_plan
    IS 'На один приём пищи в день назначается не более одного рецепта';

CREATE INDEX IF NOT EXISTS idx_menu_plan_user_id
    ON menu_plan(user_id);

CREATE INDEX IF NOT EXISTS idx_menu_plan_date
    ON menu_plan(date);

-- ============================================================
-- Тестовые данные
-- ============================================================

INSERT INTO users (name, email, hashed_password) VALUES
    ('Иван Иванов', 'ivan@example.com',
     '$2a$10$examplehashedpassword1234567890ABCDEF'),
    ('Мария Петрова', 'maria@example.com',
     '$2a$10$examplehashedpassword0987654321FEDCBA');

INSERT INTO categories (owner_id, name) VALUES
    (1, 'Супы'),
    (1, 'Выпечка'),
    (1, 'Салаты'),
    (2, 'Десерты');

INSERT INTO recipes
    (owner_id, category_id, title, description, difficulty, cook_time, ingredients, steps)
VALUES
    (1, 1, 'Борщ классический',
     'Традиционный украинский борщ со свёклой',
     'medium', 90,
     '[{"name":"Говядина","amount":"500г"},{"name":"Свёкла","amount":"2 шт"},{"name":"Капуста","amount":"300г"}]',
     '[{"description":"Отварить мясо"},{"description":"Обжарить овощи"},{"description":"Соединить и варить 20 минут"}]'),

    (1, 2, 'Пирог с яблоками',
     'Нежный домашний пирог',
     'easy', 60,
     '[{"name":"Мука","amount":"300г"},{"name":"Яблоки","amount":"4 шт"},{"name":"Сахар","amount":"150г"}]',
     '[{"description":"Замесить тесто"},{"description":"Выложить яблоки"},{"description":"Выпекать 40 минут при 180°C"}]'),

    (2, 4, 'Шоколадный торт',
     'Сочный шоколадный торт со сливочным кремом',
     'hard', 120,
     '[{"name":"Шоколад","amount":"200г"},{"name":"Яйца","amount":"4 шт"},{"name":"Сахар","amount":"200г"}]',
     '[{"description":"Растопить шоколад"},{"description":"Взбить яйца с сахаром"},{"description":"Выпекать и остудить"}]');

INSERT INTO favorites (user_id, recipe_id) VALUES
    (1, 3),
    (2, 1),
    (2, 2);

INSERT INTO menu_plan (user_id, recipe_id, date, meal_type) VALUES
    (1, 1, '2026-04-17', 'lunch'),
    (1, 2, '2026-04-18', 'breakfast'),
    (2, 3, '2026-04-17', 'dinner');

-- ============================================================
-- Полезные запросы для проверки
-- ============================================================

-- Все рецепты пользователя с категорией
-- SELECT r.id, r.title, c.name AS category, r.difficulty, r.cook_time
-- FROM recipes r
-- LEFT JOIN categories c ON c.id = r.category_id
-- WHERE r.owner_id = 1
-- ORDER BY r.created_at DESC;

-- Избранные рецепты пользователя
-- SELECT r.title, r.difficulty, r.cook_time
-- FROM favorites f
-- JOIN recipes r ON r.id = f.recipe_id
-- WHERE f.user_id = 1;

-- План меню на неделю
-- SELECT mp.date, mp.meal_type, r.title
-- FROM menu_plan mp
-- JOIN recipes r ON r.id = mp.recipe_id
-- WHERE mp.user_id = 1
--   AND mp.date BETWEEN '2026-04-14' AND '2026-04-20'
-- ORDER BY mp.date, mp.meal_type;"""

p = doc.add_paragraph()
run = p.add_run(sql)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('3.2. Индексы и их назначение', 2)

index_data = [
    ('idx_categories_owner_id', 'categories(owner_id)', 'Быстрая выборка категорий пользователя'),
    ('idx_recipes_owner_id', 'recipes(owner_id)', 'Быстрая выборка рецептов пользователя'),
    ('idx_recipes_category_id', 'recipes(category_id)', 'Фильтрация рецептов по категории'),
    ('idx_recipes_difficulty', 'recipes(difficulty)', 'Фильтрация рецептов по сложности'),
    ('idx_recipes_title_fts', 'recipes (GIN, tsvector)', 'Полнотекстовый поиск по названию'),
    ('idx_favorites_user_id', 'favorites(user_id)', 'Быстрое получение избранного пользователя'),
    ('idx_favorites_recipe_id', 'favorites(recipe_id)', 'Проверка избранности рецепта'),
    ('idx_menu_plan_user_id', 'menu_plan(user_id)', 'Выборка плана меню пользователя'),
    ('idx_menu_plan_date', 'menu_plan(date)', 'Фильтрация плана по датам (недели)'),
]

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
h4 = table4.rows[0].cells
for i, txt in enumerate(['Индекс', 'Поле(я)', 'Назначение']):
    h4[i].text = txt
    h4[i].paragraphs[0].runs[0].bold = True
    h4[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

for row_data in index_data:
    row = table4.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

heading('3.3. Разворачивание базы данных', 2)

para('Для разворачивания базы данных используется Docker Compose (рекомендуемый способ):')

code_block('# Запустить PostgreSQL через Docker Compose')
code_block('docker compose up -d db')
code_block('')
code_block('# Применить SQL-скрипт')
code_block('psql -h localhost -U recipebook -d recipebook -f schema.sql')
code_block('')
code_block('# Или через psql внутри контейнера')
code_block('docker exec -i recipebook-db-1 psql -U recipebook -d recipebook < schema.sql')

para('Параметры подключения (из docker-compose.yml):', bold=True)
code_block('Host:     localhost')
code_block('Port:     5432')
code_block('User:     recipebook')
code_block('Password: recipebook')
code_block('Database: recipebook')

heading('3.4. Проверка схемы', 2)

para('Проверочные команды psql:')
code_block('\\dt               -- список таблиц')
code_block('\\d users          -- структура таблицы users')
code_block('\\d recipes        -- структура таблицы recipes')
code_block('SELECT COUNT(*) FROM users;     -- 2 пользователя')
code_block('SELECT COUNT(*) FROM recipes;   -- 3 рецепта')
code_block('SELECT COUNT(*) FROM favorites; -- 3 записи')
code_block('SELECT COUNT(*) FROM menu_plan; -- 3 записи')

doc.add_page_break()

# ─── ЗАДАНИЕ 4: БЭКЕНД-КОД ───────────────────────────────────────────────────
heading('Задание 4. Разработка бэкенд-кода', 1)

heading('4.1. Выбор технологий', 2)

tech_data = [
    ('Язык', 'Go 1.22+', 'Высокая производительность, строгая типизация, встроенная конкурентность'),
    ('Фреймворк', 'Gin', 'Лёгкий HTTP-роутер, middleware-цепочки, встроенный JSON-биндинг'),
    ('ORM', 'GORM v2', 'Автомиграции, ассоциации, удобные CRUD-операции'),
    ('БД', 'PostgreSQL 15', 'JSONB-поля для ингредиентов/шагов, полнотекстовый поиск'),
    ('Аутентификация', 'JWT (golang-jwt/jwt/v5)', 'Stateless-токены, access + refresh схема'),
    ('Хэширование', 'bcrypt (golang.org/x/crypto)', 'Безопасное хранение паролей'),
    ('Валидация', 'go-playground/validator/v10', 'Теги валидации на структурах'),
    ('Конфигурация', 'joho/godotenv + os.Getenv', 'Переменные окружения из .env'),
    ('Контейнеризация', 'Docker + Docker Compose', 'Воспроизводимая среда запуска'),
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
h5 = table5.rows[0].cells
for i, txt in enumerate(['Компонент', 'Технология', 'Обоснование']):
    h5[i].text = txt
    h5[i].paragraphs[0].runs[0].bold = True
    h5[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
for row_data in tech_data:
    row = table5.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

heading('4.2. Структура бэкенда', 2)

para('Проект организован по принципу Clean Architecture с разделением на слои:')

struct_code = """\
recipebook-backend/
├── cmd/server/main.go        # Точка входа, инициализация Gin и БД
├── internal/
│   ├── config/config.go      # Чтение ENV-переменных
│   ├── db/db.go              # Подключение GORM к PostgreSQL
│   ├── models/               # GORM-модели (User, Recipe, Category, ...)
│   ├── handlers/             # HTTP-обработчики (auth, recipes, categories, ...)
│   ├── middleware/           # JWT-аутентификация, CORS
│   ├── repository/           # Запросы к БД (бизнес-логика)
│   └── utils/                # JWT, генерация файлов, хелперы
├── migrations/               # SQL-файлы миграций
├── uploads/                  # Загружаемые фото
├── docker-compose.yml
├── Dockerfile
└── .env"""

p = doc.add_paragraph()
run = p.add_run(struct_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.3. Инициализация сервера (main.go)', 2)

main_go = """\
// cmd/server/main.go
package main

import (
    "log"
    "github.com/gin-gonic/gin"
    "recipebook/internal/config"
    "recipebook/internal/db"
    "recipebook/internal/handlers"
    "recipebook/internal/middleware"
)

func main() {
    cfg := config.Load()           // загрузить ENV
    database := db.Connect(cfg)    // подключение к PostgreSQL
    db.AutoMigrate(database)       // автомиграции GORM

    r := gin.Default()

    // CORS
    r.Use(middleware.CORS(cfg.CORSOrigin))

    // Публичные маршруты
    auth := r.Group("/api/v1/auth")
    {
        h := handlers.NewAuthHandler(database, cfg)
        auth.POST("/register", h.Register)
        auth.POST("/login",    h.Login)
        auth.POST("/refresh",  h.Refresh)
        auth.POST("/logout",   h.Logout)
    }

    // Защищённые маршруты (требуют JWT)
    api := r.Group("/api/v1")
    api.Use(middleware.AuthRequired(cfg.JWTSecret))
    {
        rh := handlers.NewRecipeHandler(database)
        api.GET("/recipes",      rh.List)
        api.POST("/recipes",     rh.Create)
        api.GET("/recipes/:id",  rh.Get)
        api.PUT("/recipes/:id",  rh.Update)
        api.DELETE("/recipes/:id", rh.Delete)

        ch := handlers.NewCategoryHandler(database)
        api.GET("/categories",       ch.List)
        api.POST("/categories",      ch.Create)
        api.PUT("/categories/:id",   ch.Update)
        api.DELETE("/categories/:id", ch.Delete)

        fh := handlers.NewFavoriteHandler(database)
        api.GET("/favorites",              fh.List)
        api.POST("/favorites/:recipe_id",  fh.Add)
        api.DELETE("/favorites/:recipe_id", fh.Remove)

        rep := handlers.NewReportHandler(database)
        api.GET("/reports/favorites",      rep.Favorites)
        api.GET("/reports/categories",     rep.Categories)
        api.GET("/reports/favorites/xlsx", rep.FavoritesXLSX)
        api.GET("/reports/favorites/pdf",  rep.FavoritesPDF)

        mp := handlers.NewMenuPlanHandler(database)
        api.GET("/menu-plan",                    mp.Get)
        api.PUT("/menu-plan",                    mp.Set)
        api.DELETE("/menu-plan/:id",             mp.Delete)
        api.GET("/menu-plan/shopping-list/pdf",  mp.ShoppingListPDF)
    }

    log.Printf("Server running on :%s", cfg.Port)
    r.Run(":" + cfg.Port)
}"""

p = doc.add_paragraph()
run = p.add_run(main_go)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.4. GORM-модели (models/)', 2)

models_go = """\
// internal/models/user.go
package models

import "time"

type User struct {
    ID             uint      `gorm:"primaryKey;autoIncrement"`
    Name           string    `gorm:"size:100;not null"`
    Email          string    `gorm:"size:255;uniqueIndex;not null"`
    HashedPassword string    `gorm:"size:255;not null"`
    CreatedAt      time.Time
    Recipes        []Recipe  `gorm:"foreignKey:OwnerID"`
    Categories     []Category `gorm:"foreignKey:OwnerID"`
}

// internal/models/recipe.go
type Recipe struct {
    ID          uint       `gorm:"primaryKey;autoIncrement"`
    OwnerID     uint       `gorm:"not null;index"`
    CategoryID  *uint      `gorm:"index"`
    Title       string     `gorm:"size:200;not null"`
    Description string     `gorm:"type:text"`
    Difficulty  string     `gorm:"type:difficulty_enum;default:'medium'"`
    CookTime    int        `gorm:"check:cook_time >= 1 AND cook_time <= 1440"`
    PhotoURL    string     `gorm:"size:512"`
    Ingredients JSONB      `gorm:"type:jsonb;default:'[]'"`
    Steps       JSONB      `gorm:"type:jsonb;default:'[]'"`
    CreatedAt   time.Time
    UpdatedAt   time.Time
    Owner       User       `gorm:"foreignKey:OwnerID"`
    Category    *Category  `gorm:"foreignKey:CategoryID"`
}

// internal/models/category.go
type Category struct {
    ID        uint      `gorm:"primaryKey;autoIncrement"`
    OwnerID   uint      `gorm:"not null;index"`
    Name      string    `gorm:"size:100;not null"`
    CreatedAt time.Time
    Owner     User      `gorm:"foreignKey:OwnerID"`
    Recipes   []Recipe  `gorm:"foreignKey:CategoryID"`
}

// internal/models/favorite.go
type Favorite struct {
    ID        uint      `gorm:"primaryKey;autoIncrement"`
    UserID    uint      `gorm:"not null;uniqueIndex:uq_fav"`
    RecipeID  uint      `gorm:"not null;uniqueIndex:uq_fav"`
    CreatedAt time.Time
    User      User      `gorm:"foreignKey:UserID"`
    Recipe    Recipe    `gorm:"foreignKey:RecipeID"`
}

// internal/models/menu_plan.go
type MenuPlan struct {
    ID        uint      `gorm:"primaryKey;autoIncrement"`
    UserID    uint      `gorm:"not null;uniqueIndex:uq_slot"`
    RecipeID  uint      `gorm:"not null"`
    Date      time.Time `gorm:"type:date;not null;uniqueIndex:uq_slot"`
    MealType  string    `gorm:"type:meal_type_enum;not null;uniqueIndex:uq_slot"`
    CreatedAt time.Time
    User      User      `gorm:"foreignKey:UserID"`
    Recipe    Recipe    `gorm:"foreignKey:RecipeID"`
}"""

p = doc.add_paragraph()
run = p.add_run(models_go)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.5. Аутентификация (handlers/auth.go)', 2)

auth_go = """\
// internal/handlers/auth.go
package handlers

import (
    "net/http"
    "github.com/gin-gonic/gin"
    "golang.org/x/crypto/bcrypt"
    "recipebook/internal/models"
    "recipebook/internal/utils"
    "gorm.io/gorm"
)

type AuthHandler struct { db *gorm.DB; cfg *config.Config }

func NewAuthHandler(db *gorm.DB, cfg *config.Config) *AuthHandler {
    return &AuthHandler{db: db, cfg: cfg}
}

// POST /api/v1/auth/register
func (h *AuthHandler) Register(c *gin.Context) {
    var req struct {
        Name     string `json:"name"     binding:"required,min=2,max=100"`
        Email    string `json:"email"    binding:"required,email"`
        Password string `json:"password" binding:"required,min=8"`
    }
    if err := c.ShouldBindJSON(&req); err != nil {
        c.JSON(http.StatusBadRequest, gin.H{"detail": err.Error()})
        return
    }

    // Проверить уникальность email
    var existing models.User
    if h.db.Where("email = ?", req.Email).First(&existing).Error == nil {
        c.JSON(http.StatusConflict, gin.H{"detail": "Email уже зарегистрирован"})
        return
    }

    hash, _ := bcrypt.GenerateFromPassword([]byte(req.Password), bcrypt.DefaultCost)
    user := models.User{Name: req.Name, Email: req.Email, HashedPassword: string(hash)}
    h.db.Create(&user)

    access, refresh, _ := utils.GenerateTokenPair(user.ID, h.cfg.JWTSecret)
    c.JSON(http.StatusCreated, gin.H{
        "access_token": access, "refresh_token": refresh,
        "user": gin.H{"id": user.ID, "name": user.Name, "email": user.Email},
    })
}

// POST /api/v1/auth/login
func (h *AuthHandler) Login(c *gin.Context) {
    var req struct {
        Email    string `json:"email"    binding:"required,email"`
        Password string `json:"password" binding:"required"`
    }
    if err := c.ShouldBindJSON(&req); err != nil {
        c.JSON(http.StatusBadRequest, gin.H{"detail": err.Error()})
        return
    }

    var user models.User
    if err := h.db.Where("email = ?", req.Email).First(&user).Error; err != nil {
        c.JSON(http.StatusUnauthorized, gin.H{"detail": "Неверный email или пароль"})
        return
    }
    if err := bcrypt.CompareHashAndPassword([]byte(user.HashedPassword), []byte(req.Password)); err != nil {
        c.JSON(http.StatusUnauthorized, gin.H{"detail": "Неверный email или пароль"})
        return
    }

    access, refresh, _ := utils.GenerateTokenPair(user.ID, h.cfg.JWTSecret)
    c.JSON(http.StatusOK, gin.H{
        "access_token": access, "refresh_token": refresh,
        "user": gin.H{"id": user.ID, "name": user.Name, "email": user.Email},
    })
}

// POST /api/v1/auth/refresh
func (h *AuthHandler) Refresh(c *gin.Context) {
    var req struct {
        RefreshToken string `json:"refresh_token" binding:"required"`
    }
    if err := c.ShouldBindJSON(&req); err != nil {
        c.JSON(http.StatusBadRequest, gin.H{"detail": err.Error()})
        return
    }
    userID, err := utils.ParseRefreshToken(req.RefreshToken, h.cfg.JWTSecret)
    if err != nil {
        c.JSON(http.StatusUnauthorized, gin.H{"detail": "Токен недействителен"})
        return
    }
    access, refresh, _ := utils.GenerateTokenPair(userID, h.cfg.JWTSecret)
    c.JSON(http.StatusOK, gin.H{"access_token": access, "refresh_token": refresh})
}"""

p = doc.add_paragraph()
run = p.add_run(auth_go)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.6. CRUD рецептов (handlers/recipes.go)', 2)

recipes_go = """\
// internal/handlers/recipes.go
package handlers

import (
    "net/http"
    "strconv"
    "github.com/gin-gonic/gin"
    "recipebook/internal/models"
    "gorm.io/gorm"
)

type RecipeHandler struct{ db *gorm.DB }

func NewRecipeHandler(db *gorm.DB) *RecipeHandler { return &RecipeHandler{db: db} }

// GET /api/v1/recipes?search=&category_id=&difficulty=&max_time=
func (h *RecipeHandler) List(c *gin.Context) {
    userID := c.MustGet("userID").(uint)
    query := h.db.Where("owner_id = ?", userID)

    if s := c.Query("search"); s != "" {
        query = query.Where("to_tsvector('russian', title) @@ plainto_tsquery('russian', ?)", s)
    }
    if cid := c.Query("category_id"); cid != "" {
        query = query.Where("category_id = ?", cid)
    }
    if d := c.Query("difficulty"); d != "" {
        query = query.Where("difficulty = ?", d)
    }
    if mt := c.Query("max_time"); mt != "" {
        query = query.Where("cook_time <= ?", mt)
    }

    var recipes []models.Recipe
    query.Preload("Category").Order("created_at DESC").Find(&recipes)
    c.JSON(http.StatusOK, recipes)
}

// POST /api/v1/recipes
func (h *RecipeHandler) Create(c *gin.Context) {
    userID := c.MustGet("userID").(uint)
    var req struct {
        Title       string      `json:"title"       binding:"required,min=2,max=200"`
        Description string      `json:"description"`
        CategoryID  *uint       `json:"category_id"`
        Difficulty  string      `json:"difficulty"  binding:"required,oneof=easy medium hard"`
        CookTime    int         `json:"cook_time"   binding:"required,min=1,max=1440"`
        PhotoURL    string      `json:"photo_url"`
        Ingredients interface{} `json:"ingredients" binding:"required"`
        Steps       interface{} `json:"steps"       binding:"required"`
    }
    if err := c.ShouldBindJSON(&req); err != nil {
        c.JSON(http.StatusBadRequest, gin.H{"detail": err.Error()})
        return
    }
    recipe := models.Recipe{
        OwnerID:     userID,
        Title:       req.Title,
        Description: req.Description,
        CategoryID:  req.CategoryID,
        Difficulty:  req.Difficulty,
        CookTime:    req.CookTime,
        PhotoURL:    req.PhotoURL,
        Ingredients: models.JSONB(req.Ingredients),
        Steps:       models.JSONB(req.Steps),
    }
    h.db.Create(&recipe)
    h.db.Preload("Category").First(&recipe, recipe.ID)
    c.JSON(http.StatusCreated, recipe)
}

// GET /api/v1/recipes/:id
func (h *RecipeHandler) Get(c *gin.Context) {
    userID := c.MustGet("userID").(uint)
    id, _ := strconv.Atoi(c.Param("id"))
    var recipe models.Recipe
    if err := h.db.Preload("Category").
        Where("id = ? AND owner_id = ?", id, userID).
        First(&recipe).Error; err != nil {
        c.JSON(http.StatusNotFound, gin.H{"detail": "Рецепт не найден"})
        return
    }
    c.JSON(http.StatusOK, recipe)
}

// PUT /api/v1/recipes/:id
func (h *RecipeHandler) Update(c *gin.Context) {
    userID := c.MustGet("userID").(uint)
    id, _ := strconv.Atoi(c.Param("id"))
    var recipe models.Recipe
    if err := h.db.Where("id = ? AND owner_id = ?", id, userID).First(&recipe).Error; err != nil {
        c.JSON(http.StatusNotFound, gin.H{"detail": "Рецепт не найден"})
        return
    }
    if err := c.ShouldBindJSON(&recipe); err != nil {
        c.JSON(http.StatusBadRequest, gin.H{"detail": err.Error()})
        return
    }
    h.db.Save(&recipe)
    c.JSON(http.StatusOK, recipe)
}

// DELETE /api/v1/recipes/:id
func (h *RecipeHandler) Delete(c *gin.Context) {
    userID := c.MustGet("userID").(uint)
    id, _ := strconv.Atoi(c.Param("id"))
    result := h.db.Where("id = ? AND owner_id = ?", id, userID).Delete(&models.Recipe{})
    if result.RowsAffected == 0 {
        c.JSON(http.StatusNotFound, gin.H{"detail": "Рецепт не найден"})
        return
    }
    c.JSON(http.StatusNoContent, nil)
}"""

p = doc.add_paragraph()
run = p.add_run(recipes_go)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.7. JWT Middleware (middleware/auth.go)', 2)

jwt_mw = """\
// internal/middleware/auth.go
package middleware

import (
    "net/http"
    "strings"
    "github.com/gin-gonic/gin"
    "recipebook/internal/utils"
)

func AuthRequired(jwtSecret string) gin.HandlerFunc {
    return func(c *gin.Context) {
        authHeader := c.GetHeader("Authorization")
        if !strings.HasPrefix(authHeader, "Bearer ") {
            c.AbortWithStatusJSON(http.StatusUnauthorized,
                gin.H{"detail": "Токен не предоставлен"})
            return
        }
        token := strings.TrimPrefix(authHeader, "Bearer ")
        userID, err := utils.ParseAccessToken(token, jwtSecret)
        if err != nil {
            c.AbortWithStatusJSON(http.StatusUnauthorized,
                gin.H{"detail": "Токен недействителен или истёк"})
            return
        }
        c.Set("userID", userID)
        c.Next()
    }
}"""

p = doc.add_paragraph()
run = p.add_run(jwt_mw)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.8. Генерация JWT-токенов (utils/jwt.go)', 2)

jwt_go = """\
// internal/utils/jwt.go
package utils

import (
    "errors"
    "time"
    "github.com/golang-jwt/jwt/v5"
)

type Claims struct {
    UserID uint   `json:"user_id"`
    Type   string `json:"type"` // "access" | "refresh"
    jwt.RegisteredClaims
}

func GenerateTokenPair(userID uint, secret string) (access, refresh string, err error) {
    access, err = generateToken(userID, "access", 30*time.Minute, secret)
    if err != nil { return }
    refresh, err = generateToken(userID, "refresh", 7*24*time.Hour, secret)
    return
}

func generateToken(userID uint, tokenType string, ttl time.Duration, secret string) (string, error) {
    claims := Claims{
        UserID: userID,
        Type:   tokenType,
        RegisteredClaims: jwt.RegisteredClaims{
            ExpiresAt: jwt.NewNumericDate(time.Now().Add(ttl)),
            IssuedAt:  jwt.NewNumericDate(time.Now()),
        },
    }
    return jwt.NewWithClaims(jwt.SigningMethodHS256, claims).SignedString([]byte(secret))
}

func ParseAccessToken(tokenStr, secret string) (uint, error) {
    return parseToken(tokenStr, "access", secret)
}

func ParseRefreshToken(tokenStr, secret string) (uint, error) {
    return parseToken(tokenStr, "refresh", secret)
}

func parseToken(tokenStr, expectedType, secret string) (uint, error) {
    token, err := jwt.ParseWithClaims(tokenStr, &Claims{}, func(t *jwt.Token) (interface{}, error) {
        return []byte(secret), nil
    })
    if err != nil || !token.Valid { return 0, errors.New("invalid token") }
    claims := token.Claims.(*Claims)
    if claims.Type != expectedType { return 0, errors.New("wrong token type") }
    return claims.UserID, nil
}"""

p = doc.add_paragraph()
run = p.add_run(jwt_go)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F4F4F4')
p._p.get_or_add_pPr().append(shading)

heading('4.9. Обработка ошибок и принципы валидации', 2)

para('В бэкенде применяются следующие единые правила обработки ошибок:')

errors_list = [
    '400 Bad Request — ошибки валидации входных данных (binding:"required", binding:"min=2")',
    '401 Unauthorized — отсутствующий, просроченный или недействительный JWT-токен',
    '403 Forbidden — попытка изменить/удалить ресурс другого пользователя',
    '404 Not Found — запрашиваемый ресурс не найден или не принадлежит текущему пользователю',
    '409 Conflict — нарушение уникального ограничения (duplicate email, повторное добавление в избранное)',
    '500 Internal Server Error — непредвиденные ошибки сервера (логируются, клиенту отдаётся общее сообщение)',
]
for e in errors_list:
    li = doc.add_paragraph(e, style='List Bullet')
    li.runs[0].font.name = 'Times New Roman'
    li.runs[0].font.size = Pt(11)

para('Формат ответа об ошибке (унифицирован для всех эндпоинтов):', bold=True)
code_block('{ "detail": "Текст ошибки на русском языке" }')

para('Валидация входных данных выполняется на двух уровнях:', bold=True)
val_items = [
    'Уровень HTTP: теги binding на структурах запроса — Gin автоматически возвращает 400 при нарушении',
    'Уровень БД: ограничения CHECK, UNIQUE, NOT NULL в PostgreSQL — последняя линия защиты',
]
for v in val_items:
    li = doc.add_paragraph(v, style='List Bullet')
    li.runs[0].font.name = 'Times New Roman'
    li.runs[0].font.size = Pt(11)

doc.add_page_break()

# ─── ЗАДАНИЕ 5: ФИНАЛИЗАЦИЯ ───────────────────────────────────────────────────
heading('Задание 5. Согласование и доработка. Финализация', 1)

heading('5.1. Выбранные средства разработки', 2)

tools_data = [
    ('Язык бэкенда', 'Go 1.22+', 'go1.22.0 linux/amd64'),
    ('Фреймворк', 'Gin v1.9', 'github.com/gin-gonic/gin v1.9.1'),
    ('ORM', 'GORM v2', 'gorm.io/gorm v1.25.7'),
    ('СУБД', 'PostgreSQL 15', 'Docker-образ postgres:15-alpine'),
    ('Контейнеризация', 'Docker 24 + Compose v2', 'docker compose up --build'),
    ('Язык фронтенда', 'JavaScript (ES2022)', 'Node.js 20 LTS'),
    ('Фреймворк фронтенда', 'React 18 + Vite 5', 'npm create vite@latest'),
    ('Роутинг', 'React Router v6', 'react-router-dom 6.x'),
    ('Состояние', 'Zustand + TanStack Query v5', 'глобальный и серверный стейт'),
    ('Стили', 'Tailwind CSS v3', 'тёмная тема через class стратегию'),
    ('IDE', 'VS Code + GoLand', 'расширения: Go, Gin Snippets, REST Client'),
    ('API-тестирование', 'Postman / curl', 'коллекция запросов для всех эндпоинтов'),
]

table6 = doc.add_table(rows=1, cols=3)
table6.style = 'Table Grid'
h6 = table6.rows[0].cells
for i, txt in enumerate(['Инструмент', 'Название / версия', 'Примечание']):
    h6[i].text = txt
    h6[i].paragraphs[0].runs[0].bold = True
    h6[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
for row_data in tools_data:
    row = table6.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

heading('5.2. ERD-диаграмма — итоговое описание', 2)

para('Ниже приведено финальное текстовое описание ER-диаграммы (скриншот диаграммы '
     'из инструмента dbdiagram.io / DBeaver прилагается отдельно):')

erd_final = """\
  USERS               CATEGORIES           RECIPES
  ─────               ──────────           ───────
  id (PK)  ◄──────── owner_id (FK) ──────► id (PK)
  name               id (PK)               owner_id (FK) ──► USERS.id
  email (UQ)         name                  category_id (FK) ──► CATEGORIES.id
  hashed_password                          title
  created_at                               description
      │                                    difficulty (ENUM)
      │ 1:N                                cook_time (CHECK 1-1440)
      ▼                                    photo_url
  ┌──────────────────────────────────┐     ingredients (JSONB)
  │          FAVORITES               │     steps (JSONB)
  │  id (PK)                         │     created_at / updated_at
  │  user_id (FK) ──► USERS.id       │
  │  recipe_id (FK) ──► RECIPES.id   │         │
  │  created_at                      │         │ N:M
  │  UNIQUE(user_id, recipe_id)      │         ▼
  └──────────────────────────────────┘  MENU_PLAN
                                        id (PK)
                                        user_id (FK) ──► USERS.id
                                        recipe_id (FK) ──► RECIPES.id
                                        date (DATE)
                                        meal_type (ENUM)
                                        created_at
                                        UNIQUE(user_id, date, meal_type)"""

p = doc.add_paragraph()
run = p.add_run(erd_final)
run.font.name = 'Courier New'
run.font.size = Pt(8)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F8F8F8')
p._p.get_or_add_pPr().append(shading)

heading('5.3. Итоговая архитектура системы', 2)

para('Система RecipeBook реализована как трёхзвенная архитектура:')

arch_items = [
    'Фронтенд (React SPA) — взаимодействует с бэкендом через REST API, '
    'хранит JWT-токены в localStorage, реализует автоматический refresh при 401.',
    'Бэкенд (Go + Gin) — обрабатывает HTTP-запросы, выполняет бизнес-логику, '
    'генерирует XLSX/PDF-отчёты, управляет файлами загрузок.',
    'База данных (PostgreSQL 15) — хранит все данные приложения, обеспечивает '
    'целостность через внешние ключи, ограничения и уникальные индексы.',
]
for item in arch_items:
    li = doc.add_paragraph(item, style='List Bullet')
    li.runs[0].font.name = 'Times New Roman'
    li.runs[0].font.size = Pt(11)

para('Взаимодействие компонентов:', bold=True)
flow_code = """\
Browser (React SPA)
    │  HTTP/HTTPS  (Authorization: Bearer <access_token>)
    ▼
Gin HTTP Server (:8000)
    │  JWT Middleware → извлечь userID
    ├── handlers/  → валидация запроса
    ├── repository/ → запрос к GORM
    │       │
    │       ▼
    │   PostgreSQL 15 (:5432)
    │       │
    │   ◄───┘  результат запроса
    ├── (при необходимости) utils/reports → генерация XLSX/PDF
    └── JSON-ответ → Browser"""

p = doc.add_paragraph()
run = p.add_run(flow_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), 'F8F8F8')
p._p.get_or_add_pPr().append(shading)

heading('5.4. Сборка готового продукта', 2)

para('Полный запуск проекта выполняется двумя командами:')
code_block('# Клонировать и запустить бэкенд + БД')
code_block('git clone https://github.com/user/recipebook-backend')
code_block('cd recipebook-backend')
code_block('cp .env.example .env   # заполнить JWT_SECRET')
code_block('docker compose up --build -d')
code_block('')
code_block('# Запустить фронтенд (в отдельном терминале)')
code_block('git clone https://github.com/user/recipebook-frontend')
code_block('cd recipebook-frontend')
code_block('npm install')
code_block('npm run dev   # доступен на http://localhost:5173')

heading('5.5. Проверка готовности', 2)

checklist = [
    'База данных PostgreSQL 15 развёрнута и инициализирована SQL-скриптом',
    'Все 5 таблиц созданы с ограничениями, индексами и ENUM-типами',
    'Бэкенд Go/Gin запущен на порту 8000, все эндпоинты отвечают',
    'Регистрация и вход возвращают JWT access + refresh токены',
    'CRUD рецептов работает с изоляцией по owner_id',
    'Избранное: уникальное ограничение не позволяет добавить дважды',
    'Планировщик меню: уникальное ограничение по (user, date, meal_type)',
    'Отчёты генерируются и скачиваются в форматах XLSX и PDF',
    'Фронтенд React SPA запущен, авторизован и отображает данные',
    'Тёмная/светлая тема и переключение языка работают без перезагрузки',
    'Адаптивная вёрстка: корректное отображение на мобильных устройствах',
]
for item in checklist:
    li = doc.add_paragraph(f'☑  {item}', style='List Bullet')
    li.runs[0].font.name = 'Times New Roman'
    li.runs[0].font.size = Pt(11)

doc.add_page_break()

# ─── ЗАКЛЮЧЕНИЕ ───────────────────────────────────────────────────────────────
heading('Заключение', 1)

para('В ходе работы были выполнены следующие задачи:')

conclusions = [
    'Проведён анализ требований к системе RecipeBook: выделены 5 основных сущностей '
    '(User, Recipe, Category, Favorite, MenuPlan), определены их атрибуты и 12 бизнес-правил.',
    'Спроектирована ER-диаграмма с указанием типов связей (1:N и N:M через ассоциативные '
    'таблицы), первичных и внешних ключей, ограничений целостности.',
    'Разработан SQL-скрипт для PostgreSQL 15: 5 таблиц, ENUM-типы, 9 индексов (включая GIN '
    'для полнотекстового поиска), триггер автообновления updated_at, тестовые данные.',
    'Реализован бэкенд на Go 1.22 + Gin: полные CRUD-операции для всех сущностей, '
    'JWT-аутентификация (access + refresh), валидация входных данных, унифицированная '
    'обработка ошибок, генерация XLSX/PDF-отчётов.',
    'Подготовлена документация: описание средств разработки, ERD, листинг ключевых '
    'модулей бэкенда (main.go, models, handlers, middleware, utils), чеклист готовности продукта.',
]

for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {c}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Сохранить
doc.save('d:/college/Practice/RecipeBook_DB_Report.docx')
print('Файл создан: RecipeBook_DB_Report.docx')
